import streamlit as st
import uuid
from pathlib import Path
import logging
import sys

# ============================================================================
# CẤU HÌNH STREAMLIT (MUST BE FIRST)
# ============================================================================
st.set_page_config(
    page_title="EasyWord - Tạo Tài Liệu Word Thông Minh",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="collapsed",
)

# Setup logging
logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")

# Fix path to ensure 'app' module can be imported
current_dir = Path(__file__).parent
if str(current_dir) not in sys.path:
    sys.path.append(str(current_dir))

# Import app modules
try:
    from docx import Document
    from app.services.report_formatter import (
        format_uploaded_stream,
        docx_to_html,
    )
    from app.config import TEMP_DIR, CONVERTAPI_SECRET
except Exception as e:
    st.error(f"❌ Import Error: {e}")
    st.code(f"Sys Path: {sys.path}")
    st.stop()


# ============================================================================
# CSS - EASYWORD BLUE THEME (MATCHING HTML)
# ============================================================================
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800&display=swap');
    
    :root {
        --primary-blue: #2563eb;
        --primary-blue-dark: #1d4ed8;
        --primary-blue-light: #3b82f6;
        --bg-gradient-start: #f0f9ff;
        --bg-gradient-end: #e0f2fe;
        --white: #ffffff;
        --gray-50: #f9fafb;
        --gray-100: #f3f4f6;
        --gray-200: #e5e7eb;
        --gray-300: #d1d5db;
        --gray-500: #6b7280;
        --gray-700: #374151;
        --gray-900: #111827;
    }
    
    /* Hide ALL Streamlit branding */
    #MainMenu {display: none !important;}
    footer {display: none !important;}
    header[data-testid="stHeader"] {display: none !important;}
    .stDeployButton {display: none !important;}
    [data-testid="stToolbar"] {display: none !important;}
    [data-testid="stDecoration"] {display: none !important;}
    [data-testid="stStatusWidget"] {display: none !important;}
    
    html, body, [class*="css"] {
        font-family: 'Inter', -apple-system, BlinkMacSystemFont, sans-serif;
    }
    
    /* Main container - Adjust top padding to push content below custom header */
    .block-container {
        padding-top: 5rem !important; /* Adjusted for fixed header */
        padding-left: 2rem !important;
        padding-right: 2rem !important;
        max-width: 100% !important;
        background: linear-gradient(180deg, var(--bg-gradient-start) 0%, var(--bg-gradient-end) 100%);
    }
    
    /* Custom Header - Fixed at top */
    .custom-header {
        position: fixed;
        top: 0;
        left: 0;
        right: 0;
        background: rgba(255, 255, 255, 0.9);
        backdrop-filter: blur(12px);
        -webkit-backdrop-filter: blur(12px);
        box-shadow: 0 1px 2px 0 rgb(0 0 0 / 0.05);
        z-index: 1000;
        padding: 1rem 0;
    }
    
    .header-container {
        max-width: 1200px;
        margin: 0 auto;
        padding: 0 2rem;
        display: flex;
        align-items: center;
        justify-content: space-between;
    }
    
    .logo-section {
        display: flex;
        align-items: center;
        gap: 0.75rem;
        cursor: pointer;
    }
    
    .logo-img {
        height: 48px;
        width: 48px;
        border-radius: 12px;
        object-fit: cover;
    }
    
    .logo-text {
        font-size: 1.5rem;
        font-weight: 700;
        color: var(--primary-blue);
        margin: 0;
    }
    
    .header-buttons {
        display: flex;
        gap: 1rem;
    }
    
    .btn-login {
        background: transparent;
        color: var(--gray-700);
        border: 2px solid var(--gray-300);
        padding: 0.75rem 1.5rem;
        border-radius: 12px;
        font-weight: 600;
        font-size: 0.9375rem;
        cursor: pointer;
        transition: all 0.3s ease;
        text-decoration: none;
    }
    
    .btn-login:hover {
        background: var(--gray-50);
        border-color: var(--gray-500);
    }
    
    .btn-register {
        background: var(--primary-blue);
        color: white;
        border: none;
        padding: 0.75rem 1.5rem;
        border-radius: 12px;
        font-weight: 600;
        font-size: 0.9375rem;
        cursor: pointer;
        transition: all 0.3s ease;
        text-decoration: none;
        box-shadow: 0 4px 6px -1px rgb(0 0 0 / 0.1);
    }
    
    .btn-register:hover {
        background: var(--primary-blue-dark);
        transform: translateY(-2px);
        box-shadow: 0 10px 15px -3px rgb(0 0 0 / 0.1);
    }
    
    /* Add spacing for fixed header */
    .main-content {
        margin-top: 100px;
    }
    
    /* Hero Section */
    .hero-container {
        text-align: center;
        padding: 3rem 2rem;
        background: white;
        border-radius: 24px;
        margin-bottom: 2rem;
        box-shadow: 0 4px 12px rgba(0,0,0,0.05);
    }
    
    .hero-title {
        font-size: 3.5rem;
        font-weight: 800;
        color: #1e293b;
        margin-bottom: 1rem;
        line-height: 1.1;
    }
    
    .hero-subtitle {
        font-size: 1.25rem;
        color: #64748b;
        max-width: 700px;
        margin: 0 auto 2rem;
        line-height: 1.6;
    }
    
    /* Features Section (Full Width) */
    .features-section {
        background: white;
        padding: 3rem 2rem;
        border-radius: 24px;
        margin: 2rem 0;
    }
    
    /* Feature Cards */
    .feature-grid {
        display: grid;
        grid-template-columns: repeat(auto-fit, minmax(300px, 1fr));
        gap: 1.5rem;
        margin: 2rem 0;
    }
    
    .feature-grid-full {
        display: grid;
        grid-template-columns: repeat(3, 1fr);
        gap: 1.5rem;
        margin: 0 auto;
        max-width: 1200px;
    }
    
    @media (max-width: 992px) {
        .feature-grid-full {
            grid-template-columns: repeat(2, 1fr);
        }
    }
    
    @media (max-width: 600px) {
        .feature-grid-full {
            grid-template-columns: 1fr;
        }
    }
    
    .feature-card {
        background: white;
        border-radius: 20px;
        padding: 2rem;
        transition: all 0.3s ease;
        border: 2px solid transparent;
    }
    
    .feature-card:hover {
        transform: translateY(-8px);
        box-shadow: 0 12px 24px rgba(0,0,0,0.1);
        border-color: var(--primary-blue-light);
    }
    
    .feature-icon {
        width: 64px;
        height: 64px;
        border-radius: 12px;
        display: flex;
        align-items: center;
        justify-content: center;
        font-size: 2rem;
        margin-bottom: 1rem;
    }
    
    .icon-blue { background: linear-gradient(135deg, var(--primary-blue), var(--primary-blue-light)); }
    .icon-green { background: linear-gradient(135deg, #059669, #10b981); }
    .icon-purple { background: linear-gradient(135deg, #7c3aed, #8b5cf6); }
    .icon-orange { background: linear-gradient(135deg, #d97706, #f59e0b); }
    .icon-red { background: linear-gradient(135deg, #dc2626, #ef4444); }
    .icon-teal { background: linear-gradient(135deg, #0d9488, #14b8a6); }
    
    .feature-title {
        font-size: 1.25rem;
        font-weight: 700;
        color: #1e293b;
        margin-bottom: 0.5rem;
    }
    
    .feature-desc {
        font-size: 0.95rem;
        color: #64748b;
        line-height: 1.6;
    }
    
    /* Upload Card */
    .upload-card {
        background: white;
        border-radius: 20px;
        box-shadow: 0 4px 20px rgba(0,0,0,0.08);
        max-width: 600px;
        margin: 0 auto 3rem;
        overflow: hidden;
    }
    
    .upload-tabs {
        display: flex;
        border-bottom: 1px solid #e5e7eb;
    }
    
    .upload-tab {
        flex: 1;
        padding: 1rem;
        text-align: center;
        font-weight: 500;
        color: #64748b;
        cursor: pointer;
        border: none;
        background: transparent;
        transition: all 0.3s ease;
    }
    
    .upload-tab.active {
        color: var(--primary-blue);
        border-bottom: 2px solid var(--primary-blue);
        background: #f0f9ff;
    }
    
    .upload-area {
        padding: 2rem;
        text-align: center;
    }
    
    .upload-icon {
        width: 64px;
        height: 64px;
        background: #dbeafe;
        border-radius: 50%;
        display: flex;
        align-items: center;
        justify-content: center;
        margin: 0 auto 1rem;
        font-size: 1.5rem;
    }
    
    .upload-text {
        font-weight: 600;
        color: #1e293b;
        margin-bottom: 0.5rem;
    }
    
    .upload-hint {
        color: #64748b;
        font-size: 0.875rem;
        margin-bottom: 1rem;
    }
    
    .browse-btn {
        background: white;
        border: 1px solid #d1d5db;
        color: #374151;
        padding: 0.5rem 1.5rem;
        border-radius: 8px;
        font-weight: 500;
        cursor: pointer;
    }
    
    .browse-btn:hover {
        background: #f9fafb;
    }
    
    .process-btn {
        display: block;
        width: calc(100% - 2rem);
        margin: 0 1rem 1rem;
        background: var(--primary-blue);
        color: white;
        border: none;
        padding: 1rem;
        border-radius: 12px;
        font-weight: 600;
        font-size: 1rem;
        cursor: pointer;
        transition: all 0.3s ease;
    }
    
    .process-btn:hover {
        background: var(--primary-blue-dark);
    }
    
    /* CTA Section - Full Width Blue */
    .cta-section {
        background: linear-gradient(135deg, #1e3a5f 0%, #2563eb 100%);
        padding: 4rem 2rem;
        text-align: center;
        color: white;
        margin: 3rem -2rem;
        width: calc(100% + 4rem);
    }
    
    .cta-title {
        font-size: 2.5rem;
        font-weight: 700;
        margin-bottom: 1rem;
        font-style: italic;
    }
    
    .cta-subtitle {
        font-size: 1rem;
        opacity: 0.9;
        margin-bottom: 2rem;
    }
    
    .cta-btn {
        background: transparent;
        border: 2px solid white;
        color: white;
        padding: 0.875rem 2rem;
        border-radius: 12px;
        font-weight: 600;
        cursor: pointer;
        transition: all 0.3s ease;
    }
    
    .cta-btn:hover {
        background: white;
        color: var(--primary-blue);
    }
    
    /* Feature Section Background */
    .features-wrapper {
        background: #f8fafc;
        padding: 4rem 2rem;
        margin: 2rem -2rem;
        width: calc(100% + 4rem);
    }
    
    .features-title {
        text-align: center;
        font-size: 2rem;
        font-weight: 700;
        color: #1e293b;
        margin-bottom: 0.5rem;
    }
    
    .features-subtitle {
        text-align: center;
        color: #64748b;
        margin-bottom: 3rem;
    }
    
    /* Custom Footer - FULL WIDTH */
    .custom-footer {
        background: var(--gray-900);
        color: white;
        padding: 3rem 2rem 1.5rem;
        margin-top: 4rem;
        margin-left: -2rem;
        margin-right: -2rem;
        margin-bottom: -10rem;
        width: calc(100% + 4rem);
    }
    
    .footer-container {
        max-width: 1200px;
        margin: 0 auto;
        padding: 0 2rem;
    }
    
    .footer-content {
        display: grid;
        grid-template-columns: 1.5fr 1fr;
        gap: 3rem;
        margin-bottom: 2rem;
    }
    
    .footer-brand {
        max-width: 300px;
    }
    
    .footer-logo {
        height: 48px;
        width: 48px;
        border-radius: 12px;
        margin-bottom: 1rem;
    }
    
    .footer-desc {
        color: var(--gray-300);
        line-height: 1.6;
        font-size: 0.9375rem;
    }
    
    .footer-links {
        display: grid;
        grid-template-columns: repeat(3, 1fr);
        gap: 2rem;
    }
    
    .footer-column h4 {
        font-size: 0.875rem;
        font-weight: 700;
        text-transform: uppercase;
        margin-bottom: 1rem;
        color: white;
    }
    
    .footer-column a {
        display: block;
        color: var(--gray-300);
        text-decoration: none;
        margin-bottom: 0.5rem;
        font-size: 0.9375rem;
        transition: color 0.3s ease;
    }
    
    .footer-column a:hover {
        color: var(--primary-blue-light);
    }
    
    .footer-bottom {
        text-align: center;
        padding-top: 1.5rem;
        border-top: 1px solid var(--gray-700);
        color: var(--gray-500);
        font-size: 0.875rem;
    }
    
    /* Hide default Streamlit elements */
    .stDeployButton {display: none;}
    
    /* Success/Info boxes */
    .element-container .stSuccess {
        background: #dcfce7;
        border-left: 4px solid #22c55e;
        border-radius: 12px;
        padding: 1rem;
        color: #166534;
    }
    
    .element-container .stInfo {
        background: #dbeafe;
        border-left: 4px solid var(--primary-blue);
        border-radius: 12px;
        padding: 1rem;
        color: #1e40af;
    }
    
    /* Expander */
    .streamlit-expanderHeader {
        background: white;
        border-radius: 12px;
        font-weight: 600;
        color: #1e293b;
    }
    
    /* Responsive */
    @media (max-width: 768px) {
        .hero-title {
            font-size: 2.5rem;
        }
        
        .footer-content {
            grid-template-columns: 1fr;
        }
        
        .footer-links {
            grid-template-columns: 1fr;
        }
        
        .header-buttons {
            gap: 0.5rem;
        }
        
        .btn-login, .btn-register {
            padding: 0.625rem 1rem;
            font-size: 0.875rem;
        }
    }
</style>
""", unsafe_allow_html=True)

# ============================================================================
# CUSTOM HEADER (Like HTML)
# ============================================================================
import base64
import os

# Encode logo image (with fallback for cloud deployment)
logo_base64 = ""
logo_path = Path(__file__).parent / "logo.jpg"
if logo_path.exists():
    try:
        with open(logo_path, "rb") as img_file:
            logo_base64 = base64.b64encode(img_file.read()).decode()
    except:
        pass

st.markdown(f"""
<div class="custom-header">
    <div class="header-container">
        <div class="logo-section">
            <img src="data:image/jpeg;base64,{logo_base64}" alt="EasyWord Logo" class="logo-img">
            <div class="logo-text">EasyWord</div>
        </div>
        <div class="header-buttons">
            <button class="btn-login">Đăng nhập</button>
            <button class="btn-register">Đăng ký</button>
        </div>
    </div>
</div>
<div class="main-content"></div>
""", unsafe_allow_html=True)

def collect_options():
    return {
        "clean_whitespace": st.session_state.get("opt_clean", True),
        "normalize_font": st.session_state.get("opt_font", True),
        "adjust_margins": st.session_state.get("opt_margins", True),
        "indent_spacing": st.session_state.get("opt_spacing", True),
        "heading_detection": st.session_state.get("opt_heading", True),
        "format_tables": st.session_state.get("opt_tables", True),
        "insert_toc": st.session_state.get("opt_toc", True),
        "add_page_numbers": st.session_state.get("opt_page_numbers", True),
        "page_number_style": st.session_state.get("opt_page_style", "arabic"),
        "line_spacing": st.session_state.get("line_spacing", 1.3),
        "auto_numbered_heading": True,
    }

def convert_docx_to_pdf_cloud(docx_path, output_pdf_path):
    try:
        import requests
        api_secret = CONVERTAPI_SECRET
        if not api_secret:
            return None
        url = f"https://v2.convertapi.com/convert/docx/to/pdf?Secret={api_secret}&download=attachment"
        with open(docx_path, 'rb') as f:
            files = {'File': ('document.docx', f, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
            response = requests.post(url, files=files, timeout=60)
            if response.status_code == 200:
                with open(output_pdf_path, 'wb') as pdf_out:
                    pdf_out.write(response.content)
                return output_pdf_path
    except Exception as e:
        logging.warning(f"ConvertAPI failed: {e}")
    return None

def display_pdf_with_pdfjs(pdf_path):
    """Hiển thị PDF bằng PDF.js - không bị browser chặn"""
    import base64
    with open(pdf_path, "rb") as pdf_file:
        base64_pdf = base64.b64encode(pdf_file.read()).decode('utf-8')
    
    # PDF.js viewer HTML
    pdfjs_html = f'''
    <!DOCTYPE html>
    <html>
    <head>
        <script src="https://cdnjs.cloudflare.com/ajax/libs/pdf.js/3.11.174/pdf.min.js"></script>
        <style>
            * {{ margin: 0; padding: 0; box-sizing: border-box; }}
            body {{ background: #525659; font-family: Arial, sans-serif; }}
            .toolbar {{
                background: #323639;
                padding: 8px 16px;
                display: flex;
                align-items: center;
                gap: 12px;
                position: sticky;
                top: 0;
                z-index: 100;
            }}
            .toolbar button {{
                background: #4a4d50;
                border: none;
                color: white;
                padding: 6px 12px;
                border-radius: 4px;
                cursor: pointer;
                font-size: 14px;
            }}
            .toolbar button:hover {{ background: #5a5d60; }}
            .toolbar span {{ color: white; font-size: 14px; }}
            #pdf-container {{
                display: flex;
                flex-direction: column;
                align-items: center;
                padding: 20px;
                gap: 20px;
            }}
            .page-wrapper {{
                background: white;
                box-shadow: 0 4px 12px rgba(0,0,0,0.3);
                position: relative;
            }}
            .page-number {{
                position: absolute;
                bottom: -25px;
                left: 50%;
                transform: translateX(-50%);
                color: #ccc;
                font-size: 12px;
            }}
            canvas {{ display: block; }}
        </style>
    </head>
    <body>
        <div class="toolbar">
            <button onclick="zoomOut()">➖</button>
            <span id="zoom-level">100%</span>
            <button onclick="zoomIn()">➕</button>
            <span style="margin-left: 20px;">Tổng: <span id="page-count">0</span> trang</span>
        </div>
        <div id="pdf-container"></div>
        
        <script>
            pdfjsLib.GlobalWorkerOptions.workerSrc = 'https://cdnjs.cloudflare.com/ajax/libs/pdf.js/3.11.174/pdf.worker.min.js';
            
            let pdfDoc = null;
            let scale = 1.0;
            const pdfData = atob("{base64_pdf}");
            
            async function renderPDF() {{
                const loadingTask = pdfjsLib.getDocument({{data: pdfData}});
                pdfDoc = await loadingTask.promise;
                document.getElementById('page-count').textContent = pdfDoc.numPages;
                renderAllPages();
            }}
            
            async function renderAllPages() {{
                const container = document.getElementById('pdf-container');
                container.innerHTML = '';
                
                for (let i = 1; i <= pdfDoc.numPages; i++) {{
                    const page = await pdfDoc.getPage(i);
                    const viewport = page.getViewport({{ scale: scale }});
                    
                    const wrapper = document.createElement('div');
                    wrapper.className = 'page-wrapper';
                    
                    const canvas = document.createElement('canvas');
                    canvas.width = viewport.width;
                    canvas.height = viewport.height;
                    
                    const pageNum = document.createElement('div');
                    pageNum.className = 'page-number';
                    pageNum.textContent = 'Trang ' + i;
                    
                    wrapper.appendChild(canvas);
                    wrapper.appendChild(pageNum);
                    container.appendChild(wrapper);
                    
                    const ctx = canvas.getContext('2d');
                    await page.render({{ canvasContext: ctx, viewport: viewport }}).promise;
                }}
            }}
            
            function zoomIn() {{
                scale = Math.min(scale + 0.25, 3.0);
                document.getElementById('zoom-level').textContent = Math.round(scale * 100) + '%';
                renderAllPages();
            }}
            
            function zoomOut() {{
                scale = Math.max(scale - 0.25, 0.5);
                document.getElementById('zoom-level').textContent = Math.round(scale * 100) + '%';
                renderAllPages();
            }}
            
            renderPDF();
        </script>
    </body>
    </html>
    '''
    st.components.v1.html(pdfjs_html, height=800, scrolling=True)

def display_preview(doc: Document):
    temp_docx = TEMP_DIR / f"preview_{uuid.uuid4()}.docx"
    temp_pdf = TEMP_DIR / f"preview_{uuid.uuid4()}.pdf"
    try:
        doc.save(str(temp_docx))
        
        # Thử convert sang PDF để xem chính xác số trang
        if CONVERTAPI_SECRET:
            with st.spinner("🔄 Đang tạo PDF Preview..."):
                result_pdf = convert_docx_to_pdf_cloud(temp_docx, temp_pdf)
                if result_pdf and Path(result_pdf).exists():
                    st.success("✅ PDF Preview sẵn sàng!")
                    display_pdf_with_pdfjs(temp_pdf)
                    return
        
        # Fallback to HTML nếu không có API key hoặc convert thất bại
        st.info("📄 Hiển thị HTML Preview")
        html_content = docx_to_html(doc)
        st.components.v1.html(html_content, height=800, scrolling=True)
        
    except Exception as e:
        st.error(f"Lỗi Preview: {e}")
    finally:
        try:
            if temp_docx.exists(): temp_docx.unlink()
            if temp_pdf.exists(): temp_pdf.unlink()
        except: pass

# ============================================================================
# MAIN CONTENT
# ============================================================================

# Hero Section (Clean, centered)
st.markdown("""
<div style="text-align: center; padding: 2rem 0 1rem;">
    <h1 style="font-size: 3rem; font-weight: 700; color: #1e293b; margin-bottom: 1rem; line-height: 1.2;">
        Tạo Tài Liệu Word Chuyên Nghiệp<br>Trong Tích Tắc
    </h1>
    <p style="font-size: 1.125rem; color: #64748b; max-width: 600px; margin: 0 auto;">
        Upload file định dạng thô của bạn và để EasyWord xử lý mọi thứ với công nghệ AI tiên tiến. Tiết kiệm 90% thời gian định dạng.
    </p>
</div>
""", unsafe_allow_html=True)

# UPLOAD & PROCESSING SECTION
# ============================================================================

# Upload tabs (visual only - functionality via Streamlit)
st.markdown("""
<div class="upload-card" style="max-width: 600px; margin: 0 auto 1rem;">
    <div class="upload-tabs">
        <div class="upload-tab active">📤 Upload File</div>
        <div class="upload-tab">⚡ Test Nhanh</div>
    </div>
</div>
""", unsafe_allow_html=True)

# Centered upload
col_upload = st.columns([1, 4, 1])[1]
with col_upload:
    uploaded_file = st.file_uploader("Kéo thả hoặc chọn file Word (.docx)", type=["docx"], label_visibility="collapsed")
    if uploaded_file:
        st.success(f"✅ Đã chọn: **{uploaded_file.name}**")

# Options Section (Collapsible)
with st.expander("⚙️ Tùy chỉnh định dạng", expanded=False):
    col1, col2, col3 = st.columns(3)
    with col1:
        st.checkbox("🧹 Xóa dòng trống thừa", value=True, key="opt_clean")
        st.checkbox("🔤 Chuẩn hóa font chữ", value=True, key="opt_font")
        st.checkbox("📏 Thiết lập lề chuẩn", value=True, key="opt_margins")
    with col2:
        st.checkbox("↔️ Thụt đầu dòng & giãn dòng", value=True, key="opt_spacing")
        st.checkbox("🎯 Nhận diện tiêu đề", value=True, key="opt_heading")
        st.checkbox("📊 Format bảng biểu", value=True, key="opt_tables")
    with col3:
        st.checkbox("📑 Tạo mục lục", value=True, key="opt_toc")
        st.checkbox("🔢 Đánh số trang", value=True, key="opt_page_numbers")
        st.number_input("Giãn dòng", value=1.3, step=0.1, key="line_spacing")

# Process Button
col_btn = st.columns([1, 2, 1])[1]
with col_btn:
    process_clicked = st.button("✨ Bắt đầu xử lý ngay", type="primary", use_container_width=True)

# Handle file processing
if process_clicked and uploaded_file:
    with st.spinner("Đang xử lý tài liệu..."):
        try:
            file_bytes = uploaded_file.read()
            options = collect_options()
            stream, filename = format_uploaded_stream(file_bytes, uploaded_file.name, options)
            st.session_state["formatted_stream"] = stream
            st.session_state["formatted_filename"] = filename
            stream.seek(0)
            st.session_state["formatted_doc"] = Document(stream)
            st.success("✅ Chuẩn hóa thành công!")
            st.toast("Xử lý hoàn tất!", icon="🎉")
        except Exception as e:
            st.error(f"❌ Lỗi: {e}")

# Quick Test Section
with st.expander("⚡ Test Nhanh với file mẫu", expanded=False):
    col_test1, col_test2 = st.columns([3, 1])
    with col_test1:
        st.info("📁 Click nút bên cạnh để test nhanh với file `test.docx`")
    with col_test2:
        if st.button("🚀 TEST NGAY!", use_container_width=True):
            test_path = Path("test.docx")
            if test_path.exists():
                with st.spinner(f"Đang xử lý {test_path.name}..."):
                    try:
                        with open(test_path, "rb") as f:
                            file_bytes = f.read()
                        options = collect_options()
                        stream, filename = format_uploaded_stream(file_bytes, test_path.name, options)
                        st.session_state["formatted_stream"] = stream
                        st.session_state["formatted_filename"] = filename
                        stream.seek(0)
                        st.session_state["formatted_doc"] = Document(stream)
                        st.success("✅ Test file đã được xử lý thành công!")
                    except Exception as e:
                        st.error(f"❌ Lỗi: {e}")
            else:
                st.warning("⚠️ File test.docx không tồn tại.")

# ==================== RESULTS SECTION ====================
if "formatted_stream" in st.session_state:
    st.markdown("---")
    st.markdown("### 📥 Kết quả")
    
    col1, col2 = st.columns([2, 1])
    with col1:
        st.info(f"**File:** {st.session_state['formatted_filename']}")
    with col2:
        st.session_state["formatted_stream"].seek(0)
        st.download_button(
            "⬇️ Tải File Về",
            st.session_state["formatted_stream"],
            file_name=st.session_state["formatted_filename"],
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
    
    st.markdown("---")
    st.markdown("### 👁️ Xem Trước")
    with st.expander("📄 Mở Preview", expanded=True):
        if "formatted_doc" in st.session_state:
            display_preview(st.session_state["formatted_doc"])

# ============================================================================
# FEATURES SECTION
# ============================================================================
st.markdown("""
<div class="features-wrapper">
    <h2 class="features-title">EasyWord Làm Được Gì?</h2>
    <p class="features-subtitle">Khám phá các tính năng mạnh mẽ giúp công việc của bạn hiệu quả hơn</p>
</div>
""", unsafe_allow_html=True)

# Row 1 - 3 features
col1, col2, col3 = st.columns(3)
with col1:
    st.markdown("""
    <div class="feature-card">
        <div class="feature-icon icon-blue">📝</div>
        <div class="feature-title">Tự Động Định Dạng</div>
        <p class="feature-desc">AI tự động nhận diện và áp dụng định dạng chuẩn (Heading, Paragraph, List) cho tài liệu của bạn ngay lập tức.</p>
    </div>
    """, unsafe_allow_html=True)
with col2:
    st.markdown("""
    <div class="feature-card">
        <div class="feature-icon icon-green">✅</div>
        <div class="feature-title">Kiểm Tra Chính Tả</div>
        <p class="feature-desc">Phát hiện và sửa lỗi chính tả, ngữ pháp tự động với độ chính xác cao dành cho Tiếng Việt.</p>
    </div>
    """, unsafe_allow_html=True)
with col3:
    st.markdown("""
    <div class="feature-card">
        <div class="feature-icon icon-purple">📚</div>
        <div class="feature-title">Template Đa Dạng</div>
        <p class="feature-desc">Hàng trăm mẫu tài liệu chuyên nghiệp sẵn có cho mọi mục đích: Báo cáo, CV, Đơn từ, Hợp đồng.</p>
    </div>
    """, unsafe_allow_html=True)

st.markdown("")

# Row 2 - 3 features
col4, col5, col6 = st.columns(3)
with col4:
    st.markdown("""
    <div class="feature-card">
        <div class="feature-icon icon-orange">⚙️</div>
        <div class="feature-title">Tùy Chỉnh Linh Hoạt</div>
        <p class="feature-desc">Điều chỉnh mọi chi tiết theo ý muốn: font chữ, màu sắc, căn lề chỉ với vài click chuột.</p>
    </div>
    """, unsafe_allow_html=True)
with col5:
    st.markdown("""
    <div class="feature-card">
        <div class="feature-icon icon-red">⚡</div>
        <div class="feature-title">Xử Lý Siêu Nhanh</div>
        <p class="feature-desc">Xử lý tài liệu trong vài giây dù file lớn hay phức tạp. Không cần chờ đợi.</p>
    </div>
    """, unsafe_allow_html=True)
with col6:
    st.markdown("""
    <div class="feature-card">
        <div class="feature-icon icon-teal">🔒</div>
        <div class="feature-title">Bảo Mật Tuyệt Đối</div>
        <p class="feature-desc">Mọi tài liệu được mã hóa end-to-end, đảm bảo an toàn riêng tư. File tự hủy sau 24h.</p>
    </div>
    """, unsafe_allow_html=True)

# ============================================================================
# CTA SECTION
# ============================================================================
st.markdown("""
<div class="cta-section">
    <h2 class="cta-title">Sẵn Sàng Bắt Đầu?</h2>
    <p class="cta-subtitle">Tham gia hàng nghìn người dùng đang tin dùng EasyWord mỗi ngày để tối ưu hóa công việc.</p>
    <button class="cta-btn">Đăng Ký Miễn Phí Ngay</button>
</div>
""", unsafe_allow_html=True)

# ============================================================================
# CUSTOM FOOTER
# ============================================================================
st.markdown(f"""
<div class="custom-footer">
    <div class="footer-container">
        <div class="footer-content">
            <div class="footer-brand">
                <img src="data:image/jpeg;base64,{logo_base64}" alt="EasyWord" class="footer-logo">
                <p class="footer-desc">Giải pháp tạo tài liệu Word thông minh và chuyên nghiệp</p>
            </div>
            <div class="footer-links">
                <div class="footer-column">
                    <h4>Sản phẩm</h4>
                    <a href="#">Tính năng</a>
                    <a href="#">Bảng giá</a>
                    <a href="#">Templates</a>
                </div>
                <div class="footer-column">
                    <h4>Hỗ trợ</h4>
                    <a href="#">Trung tâm trợ giúp</a>
                    <a href="#">Liên hệ</a>
                    <a href="#">FAQ</a>
                </div>
                <div class="footer-column">
                    <h4>Pháp lý</h4>
                    <a href="#">Điều khoản</a>
                    <a href="#">Bảo mật</a>
                    <a href="#">Cookie</a>
                </div>
            </div>
        </div>
        <div class="footer-bottom">
            <p>&copy; 2026 EasyWord. All rights reserved.</p>
        </div>
    </div>
</div>
""", unsafe_allow_html=True)
