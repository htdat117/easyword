import logging
import re
from io import BytesIO
from html import escape

from docx import Document 
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm
from docx.text.paragraph import Paragraph

# --- CHÚ Ý: Đảm bảo các biến này tồn tại trong app.config của bạn ---
# Nếu chạy test độc lập, hãy bỏ comment dòng dưới để định nghĩa biến tạm
# BODY_FONT_SIZE = Pt(13)
# HEADING_FONT_SIZE = Pt(14)
# PAGE_NUMBER_FONT_SIZE = Pt(10)
# PARAGRAPH_INDENT = Cm(1.27)
# STANDARD_FONT = "Times New Roman"
# TOC_FONT_SIZE = Pt(13)
# UEL_MARGINS = {"top": Cm(2), "bottom": Cm(2), "left": Cm(3), "right": Cm(2)}

from app.config import (
    BODY_FONT_SIZE,
    HEADING_FONT_SIZE,
    PAGE_NUMBER_FONT_SIZE,
    PARAGRAPH_INDENT,
    STANDARD_FONT,
    TOC_FONT_SIZE,
    UEL_MARGINS,
)
from app.utils.options import merge_options
from app.services.docx_styles import (
    _copy_heading_style_to_toc,
    _ensure_caption_style,
    _ensure_east_asia_font,
    _format_toc_paragraphs,
    _set_run_format,
)
from app.services.docx_fields import (
    _add_page_number_field,
    _add_page_number_field_complex,
    _add_page_number_field_simple,
    format_page_number_run,
)

# =========================================================================
# MODULE-LEVEL COMPILED REGEX PATTERNS (Performance Optimization)
# =========================================================================
# Compile regex once at module load instead of every function call
NUMBERED_HEADING_PATTERN = re.compile(r'^(\d+(?:\.\d+)*)\.\s+(.+)$')

# Separate patterns for Table vs Figure captions
TABLE_CAPTION_PATTERN = re.compile(r'^Bảng[\s\d\.]*[:\.]?\s*(.+)$', re.IGNORECASE)
FIGURE_CAPTION_PATTERN = re.compile(r'^(Hình|Sơ đồ|Biểu đồ)[\s\d\.]*[:\.]?\s*(.+)$', re.IGNORECASE)

# Combined pattern for general caption detection (backward compatibility)
CAPTION_PATTERN = re.compile(r'^(Hình|Sơ đồ|Bảng|Biểu đồ)[\s\d\.]*[:\.]?\s+(.+)$', re.IGNORECASE)
WHITESPACE_PATTERN = re.compile(r"[ \t\u00A0]{2,}")


def _paragraph_has_image(paragraph):
    """
    Kiểm tra paragraph có chứa hình ảnh không (an toàn, không làm mất hình)
    """
    try:
        # Kiểm tra trong XML: w:drawing (hình ảnh hiện đại) hoặc w:pict (hình ảnh cũ/shape)
        has_drawing = paragraph._element.xpath('.//w:drawing')
        has_pict = paragraph._element.xpath('.//w:pict')
        return bool(has_drawing or has_pict)
    except Exception:
        return False


def _clean_leading_spaces(paragraph):
    """
    Xóa khoảng trắng đầu dòng - AN TOÀN: Chỉ xử lý text, không ảnh hưởng hình ảnh
    """
    try:
        if _paragraph_has_image(paragraph):
            return
        for run in paragraph.runs:
            if not run.text:
                continue
            cleaned = run.text.lstrip()
            run.text = cleaned
            if cleaned:
                break
    except Exception:
        pass


def _collapse_internal_spaces(paragraph):
    """
    Gộp khoảng trắng thừa - AN TOÀN: Chỉ xử lý text, không ảnh hưởng hình ảnh
    """
    try:
        if _paragraph_has_image(paragraph):
            return
        for run in paragraph.runs:
            if not run.text:
                continue
            # Use module-level compiled pattern for performance
            new_text = WHITESPACE_PATTERN.sub(" ", run.text)
            if new_text != run.text:
                run.text = new_text
    except Exception:
        pass


def _remove_paragraph(paragraph):
    try:
        parent = paragraph._element.getparent()
        if parent is not None:
            parent.remove(paragraph._element)
    except Exception:
        pass


def _looks_like_heading(text):
    text = text.strip()
    if not text or len(text) > 120:
        return False
    if text.endswith("."):
        return False
    return text.isupper()

def _detect_numbered_heading(text):
    """Detect numbered heading using cached compiled pattern."""
    text = text.strip()
    if not text:
        return None, False
    # Use module-level compiled pattern for performance
    match = NUMBERED_HEADING_PATTERN.match(text)
    if match:
        number_part = match.group(1)
        level = number_part.count('.') + 1
        level = min(max(level, 1), 6)
        return level, True
    return None, False

def _document_has_toc(doc):
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            # Check simple field
            for child in run._element:
                if child.tag.endswith("fldSimple"):
                    instr = child.get(qn("w:instr"))
                    if instr and "TOC" in instr:
                        return True
            # Check complex field
            if "TOC" in run.text and "instrText" in [c.tag.split('}')[-1] for c in run._element.iter()]:
                 return True
    return False

def _find_toc_anchor(doc):
    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name if paragraph.style else ""
        if style_name.lower().startswith("heading"):
            return paragraph
    return doc.paragraphs[0] if doc.paragraphs else None

def _insert_paragraph_after(paragraph, text=""):
    try:
        new_p = OxmlElement("w:p")
        paragraph._p.addnext(new_p)
        new_para = Paragraph(new_p, paragraph._parent)
        if text:
            new_para.add_run(text)
        return new_para
    except Exception:
        logging.warning("Cannot insert paragraph after requested position.")
        return paragraph

def _add_section_break(paragraph):
    """
    Thêm section break (ngắt trang với section mới) vào paragraph
    """
    try:
        # Lấy document từ paragraph
        doc = None
        try:
            # Thử lấy document từ part
            if hasattr(paragraph, 'part') and hasattr(paragraph.part, 'document'):
                doc = paragraph.part.document
            elif hasattr(paragraph, '_parent') and hasattr(paragraph._parent, 'part'):
                doc = paragraph._parent.part.document
        except Exception:
            pass
        
        p_pr = paragraph._p.get_or_add_pPr()
        old_sect_pr = p_pr.find(qn("w:sectPr"))
        if old_sect_pr is not None:
            p_pr.remove(old_sect_pr)
        
        sect_pr = OxmlElement("w:sectPr")
        
        # Thiết lập kiểu section break: nextPage (trang mới)
        pg_type = OxmlElement("w:type")
        pg_type.set(qn("w:val"), "nextPage")
        sect_pr.append(pg_type)
        
        # Copy margins từ section hiện tại nếu có
        if doc and hasattr(doc, 'sections') and len(doc.sections) > 0:
            try:
                current_section = doc.sections[-1]
                sect_pr_margins = current_section._sectPr.find(qn("w:pgMar"))
                if sect_pr_margins is not None:
                    new_margins = OxmlElement("w:pgMar")
                    for attr in ["top", "right", "bottom", "left", "header", "footer", "gutter"]:
                        val = sect_pr_margins.get(qn(f"w:{attr}"))
                        if val is not None:
                            new_margins.set(qn(f"w:{attr}"), val)
                    sect_pr.append(new_margins)
            except Exception as e:
                logging.debug(f"Cannot copy margins: {e}")
        
        p_pr.append(sect_pr)
        logging.info("Added section break successfully")
    except Exception as e:
        logging.warning(f"Cannot add section break: {e}")

# =========================================================================
# HÀM XỬ LÝ NHẬN DIỆN VÀ ĐÁNH SỐ CAPTION AN TOÀN
# =========================================================================
def _force_caption_font(run):
    """Force set font Times New Roman 13pt cho caption run"""
    # Set qua API
    run.font.name = STANDARD_FONT
    run.font.size = TOC_FONT_SIZE
    run.font.bold = False
    run.font.italic = True
    _ensure_east_asia_font(run)
    
    # Force set trong XML
    r_pr = run._element.get_or_add_rPr()
    
    # Xóa font cũ
    for child in list(r_pr):
        tag_name = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag_name in ["rFonts", "sz", "szCs"]:
            r_pr.remove(child)
    
    # Set font Times New Roman
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), STANDARD_FONT)
    r_fonts.set(qn("w:hAnsi"), STANDARD_FONT)
    r_fonts.set(qn("w:eastAsia"), STANDARD_FONT)
    r_fonts.set(qn("w:cs"), STANDARD_FONT)
    r_pr.insert(0, r_fonts)
    
    # Set size 13pt
    sz_half_pts = int(TOC_FONT_SIZE.pt * 2)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(sz_half_pts))
    r_pr.append(sz)
    sz_cs = OxmlElement("w:szCs")
    sz_cs.set(qn("w:val"), str(sz_half_pts))
    r_pr.append(sz_cs)

def _process_captions(doc):
    """
    Xử lý captions cho Bảng và Hình riêng biệt.
    - Bảng: đánh số Bảng 1, Bảng 2, ...
    - Hình/Sơ đồ/Biểu đồ: đánh số Hình 1, Hình 2, ...
    """
    _ensure_caption_style(doc)
    figure_count = 0  # For: Hình, Sơ đồ, Biểu đồ
    table_count = 0   # For: Bảng
    
    for paragraph in doc.paragraphs:
        has_image = _paragraph_has_image(paragraph)
        text = paragraph.text.strip()
        if not text: 
            continue
        
        # Check for Table caption first
        table_match = TABLE_CAPTION_PATTERN.match(text)
        figure_match = FIGURE_CAPTION_PATTERN.match(text)
        
        new_text = None
        
        if table_match:
            # This is a Table caption
            table_count += 1
            content = table_match.group(1).strip()
            new_text = f"Bảng {table_count}: {content}"
            paragraph.style = "UEL Figure"  # Reuse same style
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        elif figure_match:
            # This is a Figure caption (Hình, Sơ đồ, Biểu đồ)
            figure_count += 1
            content = figure_match.group(2).strip()
            new_text = f"Hình {figure_count}: {content}"
            paragraph.style = "UEL Figure"
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Apply the new text if we found a caption
        if new_text:
            if has_image:
                text_replaced = False
                for run in paragraph.runs:
                    if run.text.strip():
                        if not text_replaced:
                            run.text = new_text
                            _force_caption_font(run)
                            text_replaced = True
                        else:
                            run.text = ""
            else:
                paragraph.text = "" 
                run = paragraph.add_run(new_text)
                _force_caption_font(run)
    
    if figure_count > 0 or table_count > 0:
        logging.info(f"Processed {table_count} tables and {figure_count} figures with font {STANDARD_FONT} {TOC_FONT_SIZE.pt}pt")

# =========================================================================
# HÀM CHÈN TOC (MỤC LỤC) VÀ TOF (DANH MỤC HÌNH) - THỦ CÔNG
# =========================================================================
def _collect_headings_tables_figures_single_pass(doc):
    """
    OPTIMIZED: Thu thập headings, tables, VÀ figures trong MỘT LẦN duyệt duy nhất.
    Phân biệt Bảng vs Hình để tạo danh mục riêng biệt.
    
    Returns: (headings_list, tables_list, figures_list)
        headings_list: list of (text, level, page_estimate)
        tables_list: list of (text, page_estimate) - Bảng captions
        figures_list: list of (text, page_estimate) - Hình/Sơ đồ/Biểu đồ captions
    """
    headings = []
    tables = []
    figures = []
    page_estimate = 1
    lines_per_page = 35
    line_count = 0
    
    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name if paragraph.style else ""
        text = paragraph.text.strip()
        
        if not text:
            line_count += 1
            continue
        
        # Check if heading
        if style_name.lower().startswith("heading"):
            try:
                level = int(style_name.split()[-1]) if style_name.split()[-1].isdigit() else 1
            except:
                level = 1
            level = min(max(level, 1), 6)
            headings.append((text, level, page_estimate))
        
        # Check if caption - separate tables from figures
        if style_name in ["UEL Figure", "Caption"]:
            # Check if it's a table or figure based on text content
            if TABLE_CAPTION_PATTERN.match(text):
                tables.append((text, page_estimate))
            elif FIGURE_CAPTION_PATTERN.match(text):
                figures.append((text, page_estimate))
            else:
                # Default to figures for unrecognized captions
                figures.append((text, page_estimate))
        elif TABLE_CAPTION_PATTERN.match(text):
            tables.append((text, page_estimate))
        elif FIGURE_CAPTION_PATTERN.match(text):
            figures.append((text, page_estimate))
        
        # Estimate page number
        line_count += max(1, len(text) // 80 + 1)
        if line_count >= lines_per_page:
            page_estimate += 1
            line_count = 0
    
    return headings, tables, figures


# Legacy wrapper functions for backward compatibility
def _collect_headings_and_figures_single_pass(doc):
    """Legacy wrapper for backward compatibility."""
    headings, tables, figures = _collect_headings_tables_figures_single_pass(doc)
    # Combine tables and figures for legacy callers
    all_captions = tables + figures
    return headings, all_captions


def _collect_headings(doc):
    """Legacy wrapper - calls optimized single-pass function."""
    headings, _, _ = _collect_headings_tables_figures_single_pass(doc)
    return headings


def _collect_figures(doc):
    """Legacy wrapper - calls optimized single-pass function."""
    _, tables, figures = _collect_headings_tables_figures_single_pass(doc)
    return tables + figures


def _create_toc_entry(doc, text, level, page_num, after_para):
    """
    Tạo một entry trong mục lục với format cố định Times New Roman 13pt
    """
    # Tạo paragraph mới sau after_para
    entry_para = _insert_paragraph_after(after_para)
    
    # Set paragraph format
    fmt = entry_para.paragraph_format
    fmt.tab_stops.clear_all()
    fmt.tab_stops.add_tab_stop(Cm(16), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    fmt.line_spacing = 1.5
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(3)
    
    # Thụt lề theo level (level 1 = 0, level 2 = 0.5cm, level 3 = 1cm...)
    indent = Cm((level - 1) * 0.5)
    fmt.left_indent = indent
    fmt.first_line_indent = Pt(0)
    
    # Tạo run cho text
    run_text = entry_para.add_run(text)
    run_text.font.name = STANDARD_FONT
    run_text.font.size = TOC_FONT_SIZE
    run_text.font.bold = (level == 1)  # Bold cho heading level 1
    _ensure_east_asia_font(run_text)
    _force_run_font_xml(run_text)
    
    # Tạo tab
    run_tab = entry_para.add_run("\t")
    run_tab.font.name = STANDARD_FONT
    run_tab.font.size = TOC_FONT_SIZE
    _ensure_east_asia_font(run_tab)
    
    # Tạo run cho số trang
    run_page = entry_para.add_run(str(page_num))
    run_page.font.name = STANDARD_FONT
    run_page.font.size = TOC_FONT_SIZE
    _ensure_east_asia_font(run_page)
    _force_run_font_xml(run_page)
    
    return entry_para


def _force_run_font_xml(run):
    """
    Force set font Times New Roman 13pt trong XML level
    """
    r_pr = run._element.get_or_add_rPr()
    
    # Xóa font cũ (giữ lại bold nếu có)
    for tag in ["rFonts", "sz", "szCs"]:
        old = r_pr.find(qn(f"w:{tag}"))
        if old is not None:
            r_pr.remove(old)
    
    # Set font Times New Roman
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), STANDARD_FONT)
    r_fonts.set(qn("w:hAnsi"), STANDARD_FONT)
    r_fonts.set(qn("w:eastAsia"), STANDARD_FONT)
    r_fonts.set(qn("w:cs"), STANDARD_FONT)
    r_pr.insert(0, r_fonts)
    
    # Set size 13pt (26 half-points)
    size_half_pts = int(TOC_FONT_SIZE.pt * 2)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(size_half_pts))
    r_pr.append(sz)
    sz_cs = OxmlElement("w:szCs")
    sz_cs.set(qn("w:val"), str(size_half_pts))
    r_pr.append(sz_cs)


def _force_bold_xml(run):
    """Force bold formatting through XML to ensure it takes effect."""
    r_pr = run._element.get_or_add_rPr()
    
    # Remove existing bold element if any
    old_bold = r_pr.find(qn("w:b"))
    if old_bold is not None:
        r_pr.remove(old_bold)
    
    # Add bold element
    bold = OxmlElement("w:b")
    r_pr.append(bold)


def _insert_table_of_contents(doc, options, anchor=None):
    """
    Chèn Mục lục, Danh mục Bảng biểu, và Danh mục Hình ảnh THỦ CÔNG.
    Tạo 3 section riêng biệt với font Times New Roman 13pt.
    """
    if not options.get("insert_toc", True):
        return
    
    _copy_heading_style_to_toc(doc)
    _ensure_caption_style(doc) 
    
    # OPTIMIZED: Thu thập headings, tables, figures trong 1 lần duyệt duy nhất
    headings, tables, figures = _collect_headings_tables_figures_single_pass(doc)
    
    logging.info(f"Found {len(headings)} headings, {len(tables)} tables, {len(figures)} figures")
    
    # ==================== TẠO MỤC LỤC ====================
    first_paragraph = doc.paragraphs[0] if doc.paragraphs else None
    if first_paragraph is not None:
        toc_heading = first_paragraph.insert_paragraph_before("MỤC LỤC")
    else:
        toc_heading = doc.add_paragraph("MỤC LỤC")
    
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc_heading.paragraph_format.space_after = Pt(12)
    toc_heading.paragraph_format.space_before = Pt(0)
    for run in toc_heading.runs:
        run.font.name = STANDARD_FONT
        run.font.size = TOC_FONT_SIZE
        run.font.bold = True
        _ensure_east_asia_font(run)
        _force_run_font_xml(run)
        _force_bold_xml(run)  # Force bold through XML
    
    # Tạo các entry mục lục thủ công
    current_para = toc_heading
    
    if headings:
        for text, level, page_num in headings:
            current_para = _create_toc_entry(doc, text, level, page_num, current_para)
    else:
        # Nếu không có heading, thêm placeholder
        placeholder = _insert_paragraph_after(current_para, "(Chưa có mục lục - Hãy thêm các Heading vào văn bản)")
        placeholder.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in placeholder.runs:
            run.font.name = STANDARD_FONT
            run.font.size = TOC_FONT_SIZE
            run.font.italic = True
            _ensure_east_asia_font(run)
        current_para = placeholder
    
    # Page break sau mục lục
    toc_page_break = _insert_paragraph_after(current_para)
    toc_page_break.add_run().add_break(WD_BREAK.PAGE)
    current_para = toc_page_break
    
    # ==================== TẠO DANH MỤC BẢNG BIỂU ====================
    tot_heading = _insert_paragraph_after(current_para, "DANH MỤC BẢNG BIỂU")
    tot_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tot_heading.paragraph_format.space_before = Pt(0)
    tot_heading.paragraph_format.space_after = Pt(12)
    for run in tot_heading.runs:
        run.font.name = STANDARD_FONT
        run.font.size = TOC_FONT_SIZE
        run.font.bold = True
        _ensure_east_asia_font(run)
        _force_run_font_xml(run)
        _force_bold_xml(run)  # Force bold through XML
    
    current_para = tot_heading
    
    if tables:
        for text, page_num in tables:
            current_para = _create_toc_entry(doc, text, 1, page_num, current_para)
    else:
        # Nếu không có bảng, thêm placeholder
        placeholder = _insert_paragraph_after(current_para, "(Chưa có danh mục bảng biểu)")
        placeholder.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in placeholder.runs:
            run.font.name = STANDARD_FONT
            run.font.size = TOC_FONT_SIZE
            run.font.italic = True
            _ensure_east_asia_font(run)
        current_para = placeholder
    
    # Page break sau danh mục bảng
    tables_page_break = _insert_paragraph_after(current_para)
    tables_page_break.add_run().add_break(WD_BREAK.PAGE)
    current_para = tables_page_break
    
    # ==================== TẠO DANH MỤC HÌNH ẢNH ====================
    tof_heading = _insert_paragraph_after(current_para, "DANH MỤC HÌNH ẢNH")
    tof_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tof_heading.paragraph_format.space_before = Pt(0)
    tof_heading.paragraph_format.space_after = Pt(12)
    for run in tof_heading.runs:
        run.font.name = STANDARD_FONT
        run.font.size = TOC_FONT_SIZE
        run.font.bold = True
        _ensure_east_asia_font(run)
        _force_run_font_xml(run)
        _force_bold_xml(run)  # Force bold through XML
    
    current_para = tof_heading
    
    if figures:
        for text, page_num in figures:
            current_para = _create_toc_entry(doc, text, 1, page_num, current_para)
    else:
        # Nếu không có hình, thêm placeholder
        placeholder = _insert_paragraph_after(current_para, "(Chưa có danh mục hình ảnh)")
        placeholder.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in placeholder.runs:
            run.font.name = STANDARD_FONT
            run.font.size = TOC_FONT_SIZE
            run.font.italic = True
            _ensure_east_asia_font(run)
        current_para = placeholder
    
    # Ghi chú hướng dẫn
    hint = _insert_paragraph_after(
        current_para,
        "(* Lưu ý: Mục lục được tạo thủ công. Số trang là ước tính, vui lòng kiểm tra và chỉnh sửa nếu cần. *)",
    )
    hint.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hint.paragraph_format.space_before = Pt(12)
    for run in hint.runs:
        run.font.name = STANDARD_FONT
        run.font.size = Pt(11)
        run.font.italic = True
        run.font.color.rgb = RGBColor(128, 128, 128)
        _ensure_east_asia_font(run)
    
    # Page break cuối
    page_break_para = _insert_paragraph_after(hint)
    page_break_para.add_run().add_break(WD_BREAK.PAGE)
    
    # Tạo section break để ngắt việc đánh số trang
    _add_section_break(page_break_para)
    
    logging.info(f"Created TOC with {len(headings)} headings, {len(tables)} tables, {len(figures)} figures")

# =========================================================================
# [SỬA QUAN TRỌNG] HÀM TẠO FIELD PAGE NUMBER BẰNG COMPLEX FIELD
# =========================================================================
def _create_element(name):
    return OxmlElement(name)

def _create_attribute(element, name, value):
    element.set(qn(name), value)

def _apply_page_numbers(doc, options):
    if not options.get("add_page_numbers", True):
        logging.info("add_page_numbers=False, skipping page numbering")
        return
    
    # Xác định kiểu đánh số
    instr_main = "PAGE"
    instr_toc = "PAGE \\* ROMAN" if options.get("page_number_style") == "roman" else "PAGE"
    
    # Kiểm tra xem có Mục lục không để xác định Section bắt đầu đánh số 1
    has_toc = _document_has_toc(doc) or options.get("insert_toc", True)
    logging.info(f"Document sections: {len(doc.sections)}, has_toc: {has_toc}")
    
    # Nếu có TOC VÀ có nhiều hơn 1 section, nội dung chính ở Section 1 (bắt đầu đánh số từ 1)
    # Nếu chỉ có 1 section, bắt đầu đánh số từ section đó
    target_section_idx = 1 if (has_toc and len(doc.sections) > 1) else 0
    logging.info(f"Target section to start numbering from 1: {target_section_idx}")
    
    for idx, section in enumerate(doc.sections):
        logging.info(f"Processing section {idx}/{len(doc.sections)-1}")
        
        # --- XỬ LÝ FOOTER ---
        # Đảm bảo footer tồn tại - truy cập footer để khởi tạo nếu chưa có
        try:
            footer = section.footer
            logging.info(f"Section {idx}: Accessed footer, paragraphs: {len(footer.paragraphs)}")
        except Exception as e:
            logging.error(f"Cannot access footer of section {idx}: {e}")
            continue
        
        # Ngắt kết nối với section trước (quan trọng để restart numbering)
        try:
            section.footer.is_linked_to_previous = False
            logging.info(f"Section {idx}: Unlinked footer from previous section")
        except AttributeError as e:
            logging.warning(f"Section {idx}: Cannot unlink footer: {e}")
            pass
        
        # Xóa các đoạn văn cũ trong footer
        old_para_count = len(footer.paragraphs)
        while footer.paragraphs:
            _remove_paragraph(footer.paragraphs[0])
        logging.info(f"Section {idx}: Removed {old_para_count} old paragraphs from footer")
        
        # Nếu là trang bìa (idx=0, có mục lục VÀ có nhiều hơn 1 section) thì KHÔNG đánh số
        # Nếu chỉ có 1 section thì vẫn phải đánh số dù có TOC
        if has_toc and idx == 0 and len(doc.sections) > 1:
            logging.info(f"Section {idx}: Skipped (cover page)")
            continue
        
        # Tạo đoạn văn mới chứa số trang
        para = footer.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para_format = para.paragraph_format
        para_format.space_before = Pt(0)
        para_format.space_after = Pt(0)
        
        # Chọn kiểu số (La mã cho TOC, Số thường cho nội dung nếu cần tách biệt, ở đây dùng chung)
        # Nếu muốn TOC dùng La Mã, bạn có thể chỉnh logic tại đây. 
        # Ví dụ: nếu idx < target_section_idx dùng instr_toc
        current_instr = instr_main
        
        # Sử dụng Complex Field thay vì Simple Field vì đáng tin cậy hơn
        # Tạo run để chứa field
        run = para.add_run("")
        
        # Format font cho run TRƯỚC khi chèn field để field thừa hưởng format này
        try:
            # Format font thông qua API
            run.font.name = STANDARD_FONT
            run.font.size = PAGE_NUMBER_FONT_SIZE
            _ensure_east_asia_font(run)
            
            logging.info(f"Section {idx}: Setting font = {STANDARD_FONT}, size = {PAGE_NUMBER_FONT_SIZE.pt}pt")
            
            # Force set font trong XML để đảm bảo field thừa hưởng
            r_pr = run._element.get_or_add_rPr()
            
            # Xóa các font properties cũ
            for tag in ["rFonts", "sz", "szCs"]:
                old = r_pr.find(qn(f"w:{tag}"))
                if old is not None:
                    r_pr.remove(old)
            
            # Set mới font properties
            r_fonts = OxmlElement("w:rFonts")
            r_fonts.set(qn("w:ascii"), STANDARD_FONT)
            r_fonts.set(qn("w:hAnsi"), STANDARD_FONT)
            r_fonts.set(qn("w:eastAsia"), STANDARD_FONT)
            r_fonts.set(qn("w:cs"), STANDARD_FONT)
            r_pr.append(r_fonts)
            
            page_size_half_pts = int(PAGE_NUMBER_FONT_SIZE.pt * 2)
            logging.info(f"Section {idx}: Font size trong XML = {page_size_half_pts} half-points ({page_size_half_pts/2}pt)")
            
            sz = OxmlElement("w:sz")
            sz.set(qn("w:val"), str(page_size_half_pts))
            r_pr.append(sz)
            sz_cs = OxmlElement("w:szCs")
            sz_cs.set(qn("w:val"), str(page_size_half_pts))
            r_pr.append(sz_cs)
        except Exception as e:
            logging.error(f"Error formatting page number font: {e}")
        
        # Chèn Complex Field - đáng tin cậy hơn Simple Field
        _add_page_number_field(run, current_instr)
        
        logging.info(f"Added page number field to footer section {idx}")

        # --- XỬ LÝ SỐ TRANG BẮT ĐẦU (RESTART NUMBERING) ---
        # Chỉ reset về 1 tại section đầu tiên của phần Nội dung
        if idx == target_section_idx:
            try:
                sect_pr = section._sectPr
                pg_num_type = sect_pr.find(qn("w:pgNumType"))
                if pg_num_type is None:
                    pg_num_type = OxmlElement("w:pgNumType")
                    sect_pr.append(pg_num_type)
                
                # Bắt buộc bắt đầu từ 1
                pg_num_type.set(qn("w:start"), "1")
                pg_num_type.set(qn("w:fmt"), "decimal")
            except Exception as e:
                logging.warning(f"Error setting start page number: {e}")

# =========================================================================
# CÁC HÀM XỬ LÝ CHÍNH
# =========================================================================
def _standardize_paragraph(paragraph, options):
    style_name = paragraph.style.name if paragraph.style else ""
    if style_name in ["UEL Figure", "Caption"] or style_name.startswith("TOC"):
        return
    
    has_image = _paragraph_has_image(paragraph)
    text = paragraph.text
    
    if options.get("clean_whitespace", True) and not has_image:
        _clean_leading_spaces(paragraph)
        _collapse_internal_spaces(paragraph)
        text = paragraph.text
        
    normalized = (text or "").strip()
    
    if not normalized and not has_image:
        _remove_paragraph(paragraph)
        return
        
    if has_image and not normalized:
        if options.get("normalize_font", True):
            for run in paragraph.runs:
                if run.text:
                    _set_run_format(run, BODY_FONT_SIZE, bold=False, italic=False)
        return
    
    is_heading = False
    heading_level = None
    
    if options.get("heading_detection", True):
        if style_name.lower().startswith("heading"):
            is_heading = True
            try:
                level_str = style_name.split()[-1]
                if level_str.isdigit():
                    heading_level = int(level_str)
            except Exception:
                heading_level = 1
        elif options.get("auto_numbered_heading", True):
            detected_level, detected_heading = _detect_numbered_heading(normalized)
            if detected_heading:
                is_heading = True
                heading_level = detected_level
                try:
                    paragraph.style = f"Heading {heading_level}"
                except Exception:
                    pass
        elif _looks_like_heading(normalized):
            is_heading = True
            heading_level = 1
            try:
                paragraph.style = "Heading 1"
            except Exception:
                pass
                
    if options.get("normalize_font", True):
        target_size = HEADING_FONT_SIZE if is_heading else BODY_FONT_SIZE
        for run in paragraph.runs:
            try:
                run_has_image = (run._element.xpath('.//w:drawing') or
                                run._element.xpath('.//w:pict'))
                if run_has_image:
                    continue 
            except Exception:
                pass
            
            if is_heading:
                bold_flag = (heading_level == 1) if heading_level is not None else False
            else:
                bold_flag = bool(run.font.bold)
            italic_flag = bool(run.font.italic)
            _set_run_format(run, target_size, bold=bold_flag, italic=italic_flag)
            
    if options.get("indent_spacing", True):
        fmt = paragraph.paragraph_format
        clean_text = paragraph.text.strip()
        if is_heading:
            fmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            fmt.space_before = Pt(0)
            fmt.space_after = Pt(6)
            fmt.first_line_indent = Pt(0)
            fmt.left_indent = Pt(0)
            fmt.line_spacing = 1.5
        else:
            fmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            fmt.line_spacing = options.get("line_spacing", 1.3)
            fmt.space_before = Pt(0)
            fmt.space_after = Pt(6)
            
            # Check if paragraph has Word list/numbering format
            has_list_format = False
            try:
                p_pr = paragraph._p.get_or_add_pPr()
                # Check for numPr (numbering properties) in paragraph
                if p_pr.find(qn("w:numPr")) is not None:
                    has_list_format = True
            except:
                pass
            
            # Check if text starts with bullet-like characters
            starts_with_bullet = clean_text.startswith(("-", "+", "•", "*", "–", "—", "›", "»", "○", "●"))
            
            # Check if text starts with number followed by dot/paren (e.g., "1.", "a)", "I.")
            starts_with_number = bool(re.match(r'^[\dIVXivx]+[.)]\s', clean_text) or 
                                       re.match(r'^[a-zA-Z][.)]\s', clean_text))
            
            # Estimate if paragraph is multi-line (rough: >80 chars per line)
            chars_per_line = 80  # approximate for standard page width
            is_multi_line = len(clean_text) > chars_per_line
            
            # Indentation Logic:
            # 1. Lists/Bullets: Apply Hanging Indent (Left 0.63cm, First Line -0.63cm)
            if has_list_format or starts_with_bullet or starts_with_number:
                fmt.left_indent = Cm(0.63)
                fmt.first_line_indent = Cm(-0.63)
            
            # 2. Short single-line text (not list): No indent
            elif not is_multi_line:
                fmt.left_indent = Pt(0)
                fmt.first_line_indent = Pt(0)
                
            # 3. Standard Body Paragraphs: First line indent
            else:
                fmt.left_indent = Pt(0)
                fmt.first_line_indent = PARAGRAPH_INDENT

def apply_standard_formatting(doc: Document, options=None):
    options = merge_options(options)
    
    if options.get("adjust_margins", True):
        try:
            for section in doc.sections:
                section.top_margin = UEL_MARGINS["top"]
                section.bottom_margin = UEL_MARGINS["bottom"]
                section.left_margin = UEL_MARGINS["left"]
                section.right_margin = UEL_MARGINS["right"]
        except Exception:
            logging.warning("Cannot apply margins to document.")
            
    _process_captions(doc)
    _copy_heading_style_to_toc(doc)
    
    for paragraph in list(doc.paragraphs):
        _standardize_paragraph(paragraph, options)
        
    if options.get("format_tables", True):
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in list(cell.paragraphs):
                        _standardize_paragraph(paragraph, options)
                        
    _insert_table_of_contents(doc, options, anchor=None)
    _copy_heading_style_to_toc(doc)
    _format_toc_paragraphs(doc)
    
    # GỌI HÀM ĐÁNH SỐ TRANG SAU CÙNG
    _apply_page_numbers(doc, options)
    
    return doc

def _add_center_line(doc, text, size=HEADING_FONT_SIZE, bold=True):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_format(para.add_run(text), size, bold=bold)
    return para

def create_template_report(payload, options=None):
    student = payload.get("studentName", "Nguyễn Văn A")
    student_id = payload.get("studentId", "K2140xxxx")
    clazz = payload.get("className", "Khoa Kinh tế đối ngoại")
    title = payload.get("reportTitle", "BÁO CÁO / TIỂU LUẬN")
    year = payload.get("year", "2024-2025")
    advisor = payload.get("advisor", "GVHD: ................................")
    location = payload.get("location", "TP. Hồ Chí Minh")
    
    doc = Document()
    _add_center_line(doc, "ĐẠI HỌC QUỐC GIA TP. HỒ CHÍ MINH", Pt(14), bold=True)
    _add_center_line(doc, "TRƯỜNG ĐẠI HỌC KINH TẾ - LUẬT", Pt(14), bold=True)
    doc.add_paragraph()
    _add_center_line(doc, title.upper(), Pt(20), bold=True)
    doc.add_paragraph()
    
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_format(
        info.add_run(
            f"Sinh viên thực hiện: {student}\n"
            f"MSSV: {student_id}\n"
            f"Lớp/Khoa: {clazz}\n"
            f"{advisor}\n"
            f"Năm học: {year}"
        ),
        BODY_FONT_SIZE,
    )
    doc.add_paragraph()
    _add_center_line(doc, f"{location}, {year}", BODY_FONT_SIZE, bold=False)
    doc.add_page_break()
    
    doc.add_heading("LỜI CAM ĐOAN", level=1)
    doc.add_paragraph("Tôi cam đoan báo cáo này do tôi thực hiện, các số liệu và trích dẫn đều được ghi rõ nguồn gốc.")
    doc.add_heading("LỜI CẢM ƠN", level=1)
    doc.add_paragraph("Tập thể tác giả xin chân thành cảm ơn các thầy cô Trường Đại học Kinh tế - Luật đã hỗ trợ trong suốt quá trình thực hiện đề tài.")
    doc.add_heading("DANH MỤC TỪ VIẾT TẮT", level=1)
    doc.add_paragraph("UEL: Trường Đại học Kinh tế - Luật\nSV: Sinh viên\nGVHD: Giảng viên hướng dẫn")
    doc.add_heading("MỞ ĐẦU", level=1)
    doc.add_paragraph(payload.get("intro", "Trình bày lý do chọn đề tài, mục tiêu, phạm vi và phương pháp nghiên cứu."))
    doc.add_heading("CHƯƠNG 1. CƠ SỞ LÝ LUẬN", level=1)
    doc.add_paragraph("Mô tả cơ sở lý thuyết, tổng quan nghiên cứu.")
    doc.add_heading("CHƯƠNG 2. THỰC TRẠNG VẤN ĐỀ", level=1)
    doc.add_paragraph(payload.get("content", "Nêu hiện trạng thu thập được, số liệu minh họa và phân tích."))
    doc.add_paragraph("Hình: Biểu đồ tăng trưởng kinh tế")
    doc.add_heading("CHƯƠNG 3. GIẢI PHÁP KIẾN NGHỊ", level=1)
    doc.add_paragraph(payload.get("solution", "Đề xuất giải pháp, kiến nghị chính sách và điều kiện thực hiện."))
    doc.add_heading("KẾT LUẬN", level=1)
    doc.add_paragraph(payload.get("conclusion", "Tóm tắt kết quả đạt được và hướng nghiên cứu tiếp theo."))
    doc.add_heading("TÀI LIỆU THAM KHẢO", level=1)
    doc.add_paragraph(payload.get("references", "APA (2019). Publication Manual of the American Psychological Association (7th ed.). APA Publishing."))
    doc.add_heading("PHỤ LỤC", level=1)
    doc.add_paragraph("Đính kèm bảng biểu, hình ảnh, phiếu khảo sát (nếu có).")
    
    return apply_standard_formatting(doc, options)

def build_report_stream(doc: Document, download_name: str):
    output_stream = BytesIO()
    doc.save(output_stream)
    output_stream.seek(0)
    return output_stream, download_name

def generate_template_stream(payload):
    options = merge_options(payload.get("options"))
    doc = create_template_report(payload, options)
    return build_report_stream(doc, "bao-cao-uel.docx")

def format_uploaded_stream(file_bytes, filename, options_payload):
    options = merge_options(options_payload)
    doc = Document(BytesIO(file_bytes))
    apply_standard_formatting(doc, options)
    # Clean filename - remove trailing underscores and ensure proper extension
    base_name = filename.rsplit('.', 1)[0] if '.' in filename else filename
    base_name = base_name.strip().rstrip('_')
    safe_name = f"formatted-{base_name}.docx"
    return build_report_stream(doc, safe_name)

def docx_to_html(doc: Document) -> str:
    """
    Convert docx document to HTML for preview.
    Improved version that better represents TOC, page numbers, and formatting.
    """
    html_parts = ['<div class="docx-preview">']
    
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        style_name = paragraph.style.name if paragraph.style else ""
        
        # Skip empty paragraphs but add spacing
        if not text:
            html_parts.append('<p class="empty-line">&nbsp;</p>')
            continue
        
        # Detect paragraph type
        is_heading = style_name.lower().startswith("heading") if style_name else False
        is_caption = style_name in ["UEL Figure", "Caption"]
        is_toc = style_name.startswith("TOC") or "MỤC LỤC" in text or "DANH MỤC" in text
        is_toc_entry = style_name.startswith("TOC") and style_name != "TOC Heading"
        
        # Check for centered titles
        is_centered_title = (
            text in ["MỤC LỤC", "DANH MỤC HÌNH ẢNH"] or
            (paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER and len(text) < 50 and text.isupper())
        )
        
        # Check if this is a hint/note line
        is_hint = text.startswith("(* Lưu ý:") or text.startswith("(*")
        
        # Determine CSS class and tag
        if is_centered_title:
            css_class = "toc-title"
            tag = "h2"
        elif is_toc_entry or (is_toc and "\t" in paragraph.text):
            css_class = "toc-entry"
            tag = "p"
        elif is_heading:
            level = 1
            try:
                level = int(style_name.split()[-1]) if style_name.split()[-1].isdigit() else 1
            except:
                level = 1
            level = min(max(level, 1), 6)
            css_class = f"heading-{level}"
            tag = f"h{level}"
        elif is_caption:
            css_class = "caption"
            tag = "p"
        elif is_hint:
            css_class = "hint"
            tag = "p"
        else:
            css_class = "body-text"
            tag = "p"
        
        # Build content with proper formatting
        content_html = ""
        
        if is_toc_entry or (is_toc and "\t" in paragraph.text):
            # Format TOC entry with dots
            # Parse text and page number
            parts = text.split("\t") if "\t" in text else [text]
            if len(parts) >= 2:
                entry_text = parts[0].strip()
                page_num = parts[-1].strip()
                content_html = f'<span class="toc-text">{escape(entry_text)}</span><span class="toc-dots"></span><span class="toc-page">{escape(page_num)}</span>'
            else:
                # Just text, no page number visible
                content_html = escape(text)
        elif paragraph.runs:
            for run in paragraph.runs:
                run_text = escape(run.text) if run.text else ""
                if not run_text.strip():
                    content_html += run_text
                    continue
                
                # Build inline styles
                styles = []
                if run.bold:
                    styles.append("font-weight: bold")
                if run.italic:
                    styles.append("font-style: italic")
                if run.underline:
                    styles.append("text-decoration: underline")
                
                if styles:
                    content_html += f'<span style="{"; ".join(styles)}">{run_text}</span>'
                else:
                    content_html += run_text
        else:
            content_html = escape(text)
        
        html_parts.append(f'<{tag} class="{css_class}">{content_html}</{tag}>')
    
    # Add tables
    for table in doc.tables:
        html_parts.append('<table class="doc-table">')
        for row in table.rows:
            html_parts.append('<tr>')
            for cell in row.cells:
                cell_text = escape(cell.text.strip().replace('\n', '<br>'))
                html_parts.append(f'<td>{cell_text}</td>')
            html_parts.append('</tr>')
        html_parts.append('</table>')
    
    html_parts.append('</div>')
    
    # Generate full HTML with improved CSS
    full_html = f"""<!DOCTYPE html>
<html lang="vi">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Xem trước tài liệu</title>
    <style>
        * {{
            box-sizing: border-box;
        }}
        body {{
            margin: 0;
            padding: 20px;
            background: #e5e5e5;
            font-family: 'Times New Roman', 'Serif', Georgia, serif;
            font-size: 13pt;
            line-height: 1.5;
        }}
        .docx-preview {{
            background: white;
            max-width: 210mm;
            min-height: 297mm;
            margin: 0 auto;
            padding: 2cm 2cm 2cm 3cm;
            box-shadow: 0 4px 20px rgba(0,0,0,0.15);
        }}
        .empty-line {{
            margin: 0.3em 0;
            height: 1em;
        }}
        .toc-title {{
            text-align: center;
            font-weight: bold;
            font-size: 14pt;
            margin: 1em 0 0.8em 0;
            text-transform: uppercase;
        }}
        .toc-entry {{
            display: flex;
            align-items: baseline;
            margin: 0.3em 0;
            padding-left: 0;
            text-indent: 0;
        }}
        .toc-text {{
            flex-shrink: 0;
        }}
        .toc-dots {{
            flex-grow: 1;
            border-bottom: 1px dotted #333;
            margin: 0 8px;
            min-width: 20px;
        }}
        .toc-page {{
            flex-shrink: 0;
            text-align: right;
        }}
        .heading-1 {{
            font-size: 14pt;
            font-weight: bold;
            margin: 1em 0 0.5em 0;
        }}
        .heading-2 {{
            font-size: 13pt;
            font-weight: bold;
            margin: 0.8em 0 0.4em 0;
        }}
        .heading-3 {{
            font-size: 13pt;
            font-weight: bold;
            margin: 0.6em 0 0.3em 0;
        }}
        .body-text {{
            text-align: justify;
            text-indent: 1.27cm;
            margin: 0.3em 0;
        }}
        .caption {{
            text-align: center;
            font-style: italic;
            margin: 0.5em 0;
        }}
        .hint {{
            text-align: center;
            font-style: italic;
            color: #666;
            font-size: 11pt;
            margin: 1em 0;
        }}
        .doc-table {{
            width: 100%;
            border-collapse: collapse;
            margin: 1em 0;
        }}
        .doc-table td {{
            border: 1px solid #333;
            padding: 8px;
            vertical-align: top;
        }}
        /* Page break simulation */
        hr.page-break {{
            border: none;
            border-top: 2px dashed #ccc;
            margin: 2em 0;
        }}
    </style>
</head>
<body>
    {''.join(html_parts)}
</body>
</html>"""
    return full_html

def docx_to_html_stream(doc: Document) -> BytesIO:
    html_content = docx_to_html(doc)
    html_bytes = html_content.encode('utf-8')
    stream = BytesIO(html_bytes)
    stream.seek(0)
    return stream